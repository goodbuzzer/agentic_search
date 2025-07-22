import streamlit as st
from style import load_css
from dotenv import load_dotenv
from graph import run_graph
from langchain.callbacks import StreamlitCallbackHandler
import os
from tavily import TavilyClient
from io import BytesIO
import pandas as pd
from agents import get_conversation_chain, get_pdf_text, get_text_chunks, get_vectorstore, get_web_documents
from utils import count_pdf_pages
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime
import plotly.graph_objects as go

# Charger les variables d'environnement
load_dotenv()

# Initialiser Tavily Client
api_key = os.getenv("TAVILY_API_KEY")
client = TavilyClient(api_key=api_key)

st.set_page_config(
    page_title="MI Copilot", 
    page_icon="🤖", 
    layout="wide",
    initial_sidebar_state="expanded"
)

def page_accueil():
    load_css()
    
    # Header principal avec animation
    st.markdown("""
        <div class="main-header">
            <h1 class="main-title">MI COPILOT 🤖</h1>
            <p class="main-subtitle">Système de Veille Intelligent</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Section des fonctionnalités
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
            <div class="feature-card">
                <div class="feature-icon">🌐</div>
                <div class="feature-title">Web Search Agent</div>
                <div class="feature-description">
                    Recherchez et analysez les dernières actualités provenant de sources en ligne sélectionnées avec une précision .
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
            <div class="feature-card">
                <div class="feature-icon">📚</div>
                <div class="feature-title">Internal Docs Agent</div>
                <div class="feature-description">
                    Téléchargez et interagissez avec vos documents internes en utilisant le traitement du langage naturel avancé.
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
            <div class="feature-card">
                <div class="feature-icon">🧠</div>
                <div class="feature-title">AI Chat Assistant</div>
                <div class="feature-description">
                    Conversez avec notre IA avancée pour obtenir des insights personnalisés et des réponses précises à vos questions.
                </div>
            </div>
        """, unsafe_allow_html=True)
    
    # Métriques de performance
    st.markdown("### 📊 Statistiques en temps réel")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        data_folder = "data"
        subfolders = [f.path for f in os.scandir(data_folder) if f.is_dir()]
        total_docs = 0
        for subfolder in subfolders:
            num_docs = len([f for f in os.listdir(subfolder) if os.path.isfile(os.path.join(subfolder, f))])
            total_docs += num_docs
        st.metric("Nombre de Documents", f"{total_docs}", "12%")
    with col2:
        data_folder = "data"
        total_pages = 0
        subfolders = [f.path for f in os.scandir(data_folder) if f.is_dir()]
        for subfolder in subfolders:
            total_pages += count_pdf_pages(subfolder)
        st.metric("Nombre de pages", f"{total_pages}", "5%")
    with col3:
        st.metric("Nombre d'utilisateurs", "4", "100%")
        
    

    # Créer un histogramme pour montrer le nombre de documents dans chaque sous-dossier
    folder_names = []
    doc_counts = []

    for subfolder in subfolders:
        folder_name = os.path.basename(subfolder)
        num_docs = len([f for f in os.listdir(subfolder) if os.path.isfile(os.path.join(subfolder, f))])
        folder_names.append(folder_name)
        doc_counts.append(num_docs)


    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=folder_names,
        y=doc_counts,
        text=doc_counts,
        textposition='outside',
        marker=dict(color='skyblue', opacity=0.6)
    ))

    fig.update_layout(
        title='Répartition des documents par catégorie',
        xaxis_title='Catégorie',
        yaxis_title='Nombre de Documents',
        xaxis=dict(tickangle=45),
        template='plotly_white',
        height= 450
    )

    st.plotly_chart(fig, use_container_width=True)
        
def page_agent():
    load_css()
    
    st.markdown("""
        <div class="main-header">
            <h2 class="main-title" style="font-size: 2.5rem;">🤖 MI Copilot - Assistant IA</h2>
            <p class="main-subtitle">Posez vos questions, obtenez des réponses intelligentes</p>
        </div>
    """, unsafe_allow_html=True)

    if "messages" not in st.session_state:
        st.session_state["messages"] = [{"role": "assistant", "content": "Bonjour ! Comment puis-je vous aider aujourd'hui ? 🚀"}]

    if "conversation_chain" not in st.session_state:
        st.session_state["conversation_chain"] = run_graph

    # Container pour le chat avec style
    with st.container():
        for msg in st.session_state["messages"]:
            with st.chat_message(msg["role"]):
                st.markdown(f'<div style="padding: 0.5rem;">{msg["content"]}</div>', unsafe_allow_html=True)

    user_input = st.chat_input("💬 Tapez votre message ici...")

    if user_input:
        st.session_state["messages"].append({"role": "user", "content": user_input})
        
        with st.chat_message("user"):
            st.markdown(f'<div style="padding: 0.5rem;">{user_input}</div>', unsafe_allow_html=True)

        with st.chat_message("assistant"):
            with st.spinner("🤔 Réflexion en cours..."):
                try:
                    result = st.session_state["conversation_chain"](user_input)
                    st.session_state["messages"].append({"role": "assistant", "content": result})
                    st.markdown(f'<div style="padding: 0.5rem;">{result}</div>', unsafe_allow_html=True)

                    # Ajouter un bouton pour télécharger la réponse dans un document Word
                    doc = Document()
                    doc.add_heading('Rapport de Veille Stratégique - MI Copilot', level=1)
                    doc.add_heading('Département : Veille Stratégique', level=3)
                    # Ajouter la date d'aujourd'hui
                    today_date = datetime.now().strftime("%d/%m/%Y")
                    doc.add_paragraph(f"Date : {today_date}")
                    
                    doc.add_paragraph(result)
                    
                    # Sauvegarder le document dans un buffer
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    st.download_button(
                        label="📄 Télécharger la réponse en Word",
                        data=doc_buffer,
                        file_name="reponse_MI_Copilot.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    error_msg = f"⚠️ Une erreur s'est produite : {str(e)}"
                    st.error(error_msg)
                    st.session_state["messages"].append({"role": "assistant", "content": error_msg})

def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Résultats')
        workbook = writer.book
        worksheet = writer.sheets['Résultats']
        
        # Formatage Excel amélioré
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#4472C4',
            'font_color': 'white'
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            worksheet.set_column(col_num, col_num, 20)
    
    return output.getvalue()

def page_web_scrap():
    load_css()
    
    st.markdown("""
        <div class="main-header">
            <h2 class="main-title" style="font-size: 2.5rem;">🌐 Web Scraper Intelligent</h2>
            <p class="main-subtitle">Recherchez et analysez le contenu web en temps réel</p>
        </div>
    """, unsafe_allow_html=True)

    # Interface de recherche améliorée
    with st.container():
        col1, col2 = st.columns([3, 1])
        
        with col1:
            search_query = st.text_input(
                "🔍 Entrez votre requête de recherche :",
                placeholder="Ex: Actualités des fintech au Nigeria",
                help="Saisissez des mots-clés pour rechercher des articles pertinents"
            )
        
        with col2:
            search_depth = st.selectbox("Profondeur", ["basic", "advanced"])

    if search_query:
        with st.spinner("🔍 Recherche en cours... Veuillez patienter"):
            try:
                response = client.search(
                    query=search_query,
                    # topic = 'news',
                    country = 'cameroon',
                    search_depth=search_depth,
                    max_results=20
                )

                if response:
                # Normalisation des résultats pour éviter l'erreur
                    expected_keys = ['title', 'url', 'content']
                    cleaned_response = []

                    for result in response['results']:  # attention ici : response est sûrement un dict avec une clé 'results'
                        cleaned_result = {key: result.get(key, "") for key in expected_keys}
                        cleaned_response.append(cleaned_result)

                    df = pd.DataFrame(cleaned_response)
                    df = df.rename(columns={"title": "Titre", "url": "URL", "content": "Contenu"})

                    # Affichage du dataframe avec style
                    st.dataframe(df, use_container_width=True, height=400)

                    # Boutons d'action
                    col1, col2, col3 = st.columns([1, 1, 2])
                    
                    with col1:
                        df_xlsx = to_excel(df)
                        st.download_button(
                            label="📥 Télécharger Excel",
                            data=df_xlsx,
                            file_name=f"recherche_{search_query.replace(' ', '_')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            type="primary"
                        )
                    
                    with col2:
                        csv = df.to_csv(index=False)
                        st.download_button(
                            label="📄 Télécharger CSV",
                            data=csv,
                            file_name=f"recherche_{search_query.replace(' ', '_')}.csv",
                            mime="text/csv"
                        )
                else:
                    st.warning("🔍 Aucun résultat trouvé pour cette recherche.")

            except Exception as e:
                st.error(f"❌ Erreur lors de la recherche : {e}")
            
def page_docs_ai():
    load_css()
    
    st.markdown("""
        <div class="main-header">
            <h2 class="main-title" style="font-size: 2.5rem;">📄 Docs AI - Analyseur Intelligent</h2>
            <p class="main-subtitle">Téléchargez vos documents et posez vos questions</p>
        </div>
    """, unsafe_allow_html=True)

    # Zone d'upload stylée
    st.markdown("""
        <div class="upload-area">
            <h3>📂 Zone de téléchargement</h3>
            <p>Glissez-déposez vos fichiers PDF ou utilisez le bouton ci-dessous</p>
        </div>
    """, unsafe_allow_html=True)
    
    uploaded_files = st.file_uploader(
        "Téléchargez des fichiers PDF", 
        type=['pdf'], 
        accept_multiple_files=True,
        help="Vous pouvez télécharger plusieurs fichiers PDF simultanément"
    )
    
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} fichier(s) téléchargé(s) avec succès !")
        
        for file in uploaded_files:
            with st.expander(f"📄 {file.name}"):
                st.write(f"**Taille :** {file.size / 1024:.2f} KB")
                st.write(f"**Type :** {file.type}")
        
        with st.spinner("Processing PDF documents..."):
                    all_chunks = []
                    for uploaded_file in uploaded_files:
                        docs = get_pdf_text(uploaded_file)  # Lire chaque PDF
                        
                        # Si c'est une liste de documents (plusieurs pages)
                        if isinstance(docs, list):
                            for doc in docs:
                                if doc["page_content"]:
                                    chunks = get_text_chunks(doc)
                                    all_chunks.extend(chunks)
                        # Si c'est un seul document
                        elif isinstance(docs, dict) and docs["page_content"]:
                            chunks = get_text_chunks(docs)
                            all_chunks.extend(chunks)
                    
                    if all_chunks:
                        st.info(f"Extracted {len(all_chunks)} text chunks from {len(uploaded_files)} documents")
                        
                        with st.spinner("Création des embeddings vectoriels..."):
                            vectorstore = get_vectorstore(all_chunks)
                            
                            if vectorstore:
                                conversation_chain = get_conversation_chain(vectorstore)
                                st.session_state["conversation_chain"] = conversation_chain  # Stockage dans la session
                                st.success(f"Document traité avec succès ! Prêt à répondre aux questions.")
                            else:
                                st.error("Failed to create embeddings. Please check your OpenAI API key.")
                    else:
                        st.warning("No readable content could be extracted from the uploaded documents.")
    
    
    if "messages" not in st.session_state:
        st.session_state["messages"] = [{"role": "assistant", "content": "Bonjour ! Comment puis-je vous aider aujourd'hui ? 🚀"}]

    if "conversation_chain" not in st.session_state:
        st.session_state["conversation_chain"] = run_graph

    # Container pour le chat avec style
    with st.container():
        for msg in st.session_state["messages"]:
            with st.chat_message(msg["role"]):
                st.markdown(f'<div style="padding: 0.5rem;">{msg["content"]}</div>', unsafe_allow_html=True)

    user_query = st.chat_input("💬 Tapez votre message ici...")

    if user_query:
        st.session_state.messages.append({"role": "user", "content": user_query})
        st.chat_message("user").write(user_query)

        if "conversation_chain" in st.session_state:
            conversation_chain = st.session_state["conversation_chain"]
            with st.chat_message("assistant"):
                st_cb = StreamlitCallbackHandler(st.container())
                try:
                    response = conversation_chain.invoke({"question": user_query})
                    answer = response["answer"]
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                    st.write(answer)
                except Exception as e:
                    error_msg = f"An error occurred: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
        else:
            warning_msg = "Please upload documents or fetch news first!"
            st.warning(warning_msg)
            st.session_state.messages.append({"role": "assistant", "content": warning_msg})        
                
        
                
    

def page_aide():
    load_css()
    
    st.markdown("""
        <div class="main-header">
            <h2 class="main-title" style="font-size: 2.5rem;">📖 Centre d'Aide</h2>
            <p class="main-subtitle">Tout ce que vous devez savoir sur MI Copilot</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Guide d'utilisation avec accordéon
    with st.expander("🚀 Guide de démarrage rapide", expanded=True):
        st.markdown("""
        ### Comment utiliser MI Copilot ?
        
        1. **🏠 Accueil** : Vue d'ensemble des fonctionnalités
        2. **🤖 Agent IA** : Chat intelligent avec l'IA
        3. **🌐 Web Scraper** : Recherche et analyse web
        4. **📄 Docs AI** : Analyse de documents PDF
        """)
    
    with st.expander("💡 Conseils d'utilisation"):
        st.markdown("""
        - Utilisez des requêtes précises pour de meilleurs résultats
        - Les documents PDF de moins de 10MB sont recommandés
        - Explorez les différents modes de recherche web
        """)
    
    with st.expander("❓ FAQ"):
        st.markdown("""
        **Q: Quels formats de fichiers sont supportés ?**
        A: Actuellement, nous supportons les fichiers PDF et textes.
        
        **Q: Y a-t-il une limite au nombre de recherches ?**
        A: Non, vous pouvez effectuer autant de recherches que nécessaire.
        """)

def main():
    load_css()
    
    # Sidebar améliorée
    with st.sidebar:
        # Logo avec style
        try:
            st.image("mi_copilot.png", width = 250)
        except:
            st.markdown("### 🤖 MI Copilot")
        
        st.markdown("---")
        
        # Navigation avec icônes
        st.markdown("### 🧭 Navigation")
        page = st.selectbox(
            "Choisissez une page :", 
            ["🏠 Accueil", "🤖 Agent IA", "🌐 Web Scraper", "📄 Docs AI", "📖 Centre d'Aide"],
            label_visibility="collapsed"
        )
        
        st.markdown("---")
        
        # Actions rapides
        st.markdown("### ⚡ Actions Rapides")
        
        if st.button("🧹 Effacer l'historique", type="secondary"):
            st.session_state["messages"] = [{"role": "assistant", "content": "Bonjour ! Comment puis-je vous aider aujourd'hui ? 🚀"}]
            st.success("Historique effacé !")
        
        if st.button("🔄 Actualiser", type="secondary"):
            st.rerun()
        
        st.markdown("---")
        
        # Informations système
        st.markdown("### ℹ️ Informations")
        st.info("**Statut :** 🟢 Actif")

    # Navigation entre les pages
    if page == "🏠 Accueil":
        page_accueil()
    elif page == "🤖 Agent IA":
        page_agent()
    elif page == "🌐 Web Scraper":
        page_web_scrap()
    elif page == "📄 Docs AI":
        page_docs_ai()
    elif page == "📖 Centre d'Aide":
        page_aide()
    
    # Footer
    st.markdown("""
        <div class="footer">
            <p>© 2025 MI Copilot - MEDIA INTELLIGENCE SARL 🚀</p>
        </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
